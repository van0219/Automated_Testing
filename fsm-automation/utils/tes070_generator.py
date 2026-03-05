"""
TES-070 Document Generator

Generates TES-070 test results documents from PyTest/Allure results.
"""

import json
import argparse
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


class TES070Generator:
    """Generate TES-070 Word documents from test results"""
    
    def __init__(self, allure_results_dir: str, output_dir: str = "reports/TES070"):
        self.allure_results_dir = Path(allure_results_dir)
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)
    
    def generate(self, interface_id: str, interface_name: str):
        """Generate TES-070 document for an interface"""
        # Load test results
        results = self._load_test_results()
        
        # Create Word document
        doc = Document()
        
        # Add title page
        self._add_title_page(doc, interface_id, interface_name)
        
        # Add test summary
        self._add_test_summary(doc, results)
        
        # Add test scenarios
        self._add_test_scenarios(doc, results)
        
        # Save document
        output_file = self.output_dir / f"{interface_id}_TES070_{datetime.now().strftime('%Y%m%d')}.docx"
        doc.save(output_file)
        
        print(f"TES-070 document generated: {output_file}")
        return output_file
    
    def _load_test_results(self) -> dict:
        """Load test results from Allure JSON files"""
        results = {
            "total": 0,
            "passed": 0,
            "failed": 0,
            "scenarios": []
        }
        
        # Find all test result JSON files
        for result_file in self.allure_results_dir.glob("*-result.json"):
            with open(result_file, 'r') as f:
                test_data = json.load(f)
                
                results["total"] += 1
                
                if test_data.get("status") == "passed":
                    results["passed"] += 1
                else:
                    results["failed"] += 1
                
                results["scenarios"].append({
                    "name": test_data.get("name", "Unknown"),
                    "status": test_data.get("status", "unknown"),
                    "description": test_data.get("description", ""),
                    "steps": test_data.get("steps", []),
                    "attachments": test_data.get("attachments", [])
                })
        
        return results
    
    def _add_title_page(self, doc: Document, interface_id: str, interface_name: str):
        """Add TES-070 title page"""
        # Title
        title = doc.add_heading(f"TES-070 Test Results Document", level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Interface ID
        interface_heading = doc.add_heading(f"{interface_id}", level=2)
        interface_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Interface Name
        name_para = doc.add_paragraph(interface_name)
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Date
        date_para = doc.add_paragraph(f"Test Date: {datetime.now().strftime('%B %d, %Y')}")
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Page break
        doc.add_page_break()
    
    def _add_test_summary(self, doc: Document, results: dict):
        """Add test execution summary"""
        doc.add_heading("Test Execution Summary", level=1)
        
        # Summary table
        table = doc.add_table(rows=4, cols=2)
        table.style = 'Light Grid Accent 1'
        
        # Headers
        table.rows[0].cells[0].text = "Metric"
        table.rows[0].cells[1].text = "Value"
        
        # Data
        table.rows[1].cells[0].text = "Total Scenarios"
        table.rows[1].cells[1].text = str(results["total"])
        
        table.rows[2].cells[0].text = "Passed"
        table.rows[2].cells[1].text = str(results["passed"])
        
        table.rows[3].cells[0].text = "Failed"
        table.rows[3].cells[1].text = str(results["failed"])
        
        doc.add_paragraph()
    
    def _add_test_scenarios(self, doc: Document, results: dict):
        """Add detailed test scenario results"""
        doc.add_heading("Test Scenario Results", level=1)
        
        for idx, scenario in enumerate(results["scenarios"], 1):
            # Scenario heading
            status_icon = "✅" if scenario["status"] == "passed" else "❌"
            doc.add_heading(f"{status_icon} Scenario {idx}: {scenario['name']}", level=2)
            
            # Description
            if scenario["description"]:
                doc.add_paragraph(scenario["description"])
            
            # Status
            status_para = doc.add_paragraph()
            status_para.add_run("Status: ").bold = True
            status_run = status_para.add_run(scenario["status"].upper())
            if scenario["status"] == "passed":
                status_run.font.color.rgb = RGBColor(0, 128, 0)
            else:
                status_run.font.color.rgb = RGBColor(255, 0, 0)
            
            # Test steps
            if scenario["steps"]:
                doc.add_heading("Test Steps:", level=3)
                for step_idx, step in enumerate(scenario["steps"], 1):
                    step_para = doc.add_paragraph(style='List Number')
                    step_para.add_run(step.get("name", "Unknown step"))
            
            # Screenshots
            if scenario["attachments"]:
                doc.add_heading("Evidence:", level=3)
                for attachment in scenario["attachments"]:
                    if attachment.get("type") == "image/png":
                        # Add screenshot
                        screenshot_path = self.allure_results_dir / attachment.get("source", "")
                        if screenshot_path.exists():
                            doc.add_paragraph(attachment.get("name", "Screenshot"))
                            doc.add_picture(str(screenshot_path), width=Inches(6))
            
            doc.add_page_break()


def main():
    """CLI entry point"""
    parser = argparse.ArgumentParser(description="Generate TES-070 documents from test results")
    parser.add_argument("--results-dir", default="reports/allure-results", help="Allure results directory")
    parser.add_argument("--output-dir", default="reports/TES070", help="Output directory for TES-070 documents")
    parser.add_argument("--interface-id", required=True, help="Interface ID (e.g., EXT_FIN_004)")
    parser.add_argument("--interface-name", required=True, help="Interface name")
    
    args = parser.parse_args()
    
    generator = TES070Generator(args.results_dir, args.output_dir)
    generator.generate(args.interface_id, args.interface_name)


if __name__ == "__main__":
    main()
