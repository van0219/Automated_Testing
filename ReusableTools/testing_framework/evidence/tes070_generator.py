"""TES-070 document generator"""

from pathlib import Path
from datetime import datetime
from typing import Dict, Any, Optional
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from ..engine.results import TestResult
from ..utils.logger import Logger


class TES070Generator:
    """
    Generate TES-070 Word documents from test results.
    
    Creates formatted test evidence documents with screenshots.
    """
    
    def __init__(self, logger: Optional[Logger] = None):
        """
        Initialize TES-070 generator.
        
        Args:
            logger: Optional logger instance
        """
        self.logger = logger
    
    def generate(
        self,
        test_result: TestResult,
        screenshot_dir: str,
        output_dir: str = "TES-070/Generated_TES070s"
    ) -> str:
        """
        Generate TES-070 document.
        
        Args:
            test_result: TestResult with scenario results
            screenshot_dir: Directory containing screenshots
            output_dir: Output directory for document
        
        Returns:
            Path to generated .docx file
        """
        if self.logger:
            self.logger.info(f"Generating TES-070 document for {test_result.interface_id}")
        
        # Create document
        doc = Document()
        
        # Add title page
        self._add_title_page(doc, test_result)
        
        # Add page break
        doc.add_page_break()
        
        # Add table of contents placeholder
        self._add_toc_placeholder(doc)
        
        # Add page break
        doc.add_page_break()
        
        # Add test summary
        self._add_test_summary(doc, test_result)
        
        # Add scenarios
        for scenario_result in test_result.scenario_results:
            self._add_scenario(doc, scenario_result, screenshot_dir)
        
        # Save document
        output_path = self._save_document(doc, test_result.interface_id, output_dir)
        
        if self.logger:
            self.logger.info(f"TES-070 document generated: {output_path}")
        
        return output_path
    
    def _add_title_page(self, doc: Document, test_result: TestResult) -> None:
        """Add title page to document."""
        # Title
        title = doc.add_heading(f"TES-070 Test Results Document", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Interface ID and Name
        doc.add_paragraph()
        p = doc.add_paragraph(f"Interface ID: {test_result.interface_id}")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Date
        doc.add_paragraph()
        date_str = datetime.now().strftime("%Y-%m-%d")
        p = doc.add_paragraph(f"Test Date: {date_str}")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    def _add_toc_placeholder(self, doc: Document) -> None:
        """Add table of contents placeholder."""
        doc.add_heading("Table of Contents", level=1)
        doc.add_paragraph(
            "Note: Press F9 in Microsoft Word to update the table of contents."
        )
    
    def _add_test_summary(self, doc: Document, test_result: TestResult) -> None:
        """Add test summary section."""
        doc.add_heading("Test Summary", level=1)
        
        total_scenarios = len(test_result.scenario_results)
        passed_scenarios = sum(1 for s in test_result.scenario_results if s.passed)
        failed_scenarios = total_scenarios - passed_scenarios
        
        doc.add_paragraph(f"Total Scenarios: {total_scenarios}")
        doc.add_paragraph(f"Passed: {passed_scenarios}")
        doc.add_paragraph(f"Failed: {failed_scenarios}")
        doc.add_paragraph(f"Overall Status: {'PASSED' if test_result.passed else 'FAILED'}")
    
    def _add_scenario(
        self,
        doc: Document,
        scenario_result: Any,
        screenshot_dir: str
    ) -> None:
        """Add scenario section with test steps."""
        # Scenario heading
        doc.add_heading(f"Scenario {scenario_result.scenario_id}: {scenario_result.title}", level=2)
        
        # Test steps table
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Grid Table 4 - Accent 1'
        
        # Header row
        header_cells = table.rows[0].cells
        header_cells[0].text = "Step"
        header_cells[1].text = "Description"
        header_cells[2].text = "Expected"
        header_cells[3].text = "Actual"
        header_cells[4].text = "Status"
        
        # Add step rows
        for step_result in scenario_result.step_results:
            row_cells = table.add_row().cells
            row_cells[0].text = str(step_result.step_number)
            row_cells[1].text = step_result.description
            row_cells[2].text = "Success"
            row_cells[3].text = step_result.action_result.message
            row_cells[4].text = "PASS" if step_result.passed else "FAIL"
            
            # Add screenshot if available
            if step_result.screenshot_path:
                doc.add_paragraph()
                doc.add_paragraph(f"Evidence - Step {step_result.step_number}:")
                try:
                    screenshot_path = Path(step_result.screenshot_path)
                    if screenshot_path.exists():
                        doc.add_picture(str(screenshot_path), width=Inches(6))
                except Exception as e:
                    if self.logger:
                        self.logger.warning(f"Failed to add screenshot: {str(e)}")
                    doc.add_paragraph(f"[Screenshot not available: {step_result.screenshot_path}]")
    
    def _save_document(self, doc: Document, interface_id: str, output_dir: str) -> str:
        """Save document to file."""
        # Create output directory
        output_path = Path(output_dir)
        output_path.mkdir(parents=True, exist_ok=True)
        
        # Generate filename
        timestamp = datetime.now().strftime("%Y%m%d")
        filename = f"{interface_id}_TES-070_{timestamp}.docx"
        filepath = output_path / filename
        
        # Save document
        doc.save(str(filepath))
        
        return str(filepath)
