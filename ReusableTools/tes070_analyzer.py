"""
TES-070 Document Analyzer - Robust analysis of test results documents

Purpose: Analyze TES-070 documents including text and tables
Use when: You need to understand test evidence format and content

Features:
- Sequential content extraction preserving document order
- Robust error handling for various document structures
- Reusable across all TES-070 document formats
"""

from docx import Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import Table
from docx.text.paragraph import Paragraph
import os
from pathlib import Path
from typing import Dict, List, Optional, Any
import json
import re
from dataclasses import dataclass, asdict
from enum import Enum


class ContentType(Enum):
    """Types of content in document"""
    TEXT = "text"
    TABLE = "table"


@dataclass
class TestStep:
    """Single test step"""
    number: str
    description: str
    result: str


@dataclass
class TestScenario:
    """Complete test scenario with evidence"""
    title: str
    description: str
    test_steps: List[TestStep]
    results: List[str]
    content_flow: List[Dict[str, Any]]


class TableParser:
    """Handles table parsing and structure detection"""
    
    @staticmethod
    def parse(table: Table) -> Dict[str, Any]:
        """Parse table data with structure"""
        try:
            rows = []
            
            for row in table.rows:
                cells = []
                for cell in row.cells:
                    cells.append(cell.text.strip())
                rows.append(cells)
            
            return {
                'rows': rows,
                'row_count': len(rows),
                'col_count': len(rows[0]) if rows else 0
            }
        except Exception as e:
            return {
                'rows': [],
                'row_count': 0,
                'col_count': 0,
                'error': str(e)
            }
    
    @staticmethod
    def is_test_steps_table(table_data: Dict[str, Any]) -> bool:
        """Check if table contains test steps"""
        try:
            if not table_data.get('rows') or len(table_data['rows']) < 2:
                return False
            
            header = table_data['rows'][0]
            
            # Check for step-related columns (No., Step, Steps, #, etc.)
            has_step_column = any(keyword in cell.lower() for cell in header 
                                 for keyword in ['step', 'no.', 'no', '#'])
            
            # Check for result column (can be empty header)
            has_result_column = (
                any('result' in cell.lower() for cell in header) or
                len(header) >= 3  # 3+ columns often means: No, Steps, Result
            )
            
            # Additional validation: check if first data row looks like a test step
            if len(table_data['rows']) > 1:
                first_row = table_data['rows'][1]
                # First cell should be a number or step indicator
                if first_row and (first_row[0].strip().isdigit() or 
                                 first_row[0].strip() in ['1', '2', '3', '4', '5']):
                    return True
            
            return has_step_column and has_result_column
        except Exception:
            return False
    
    @staticmethod
    def parse_test_steps(table_data: Dict[str, Any]) -> List[TestStep]:
        """Parse test steps from table"""
        steps = []
        
        try:
            # Skip header row
            for row in table_data['rows'][1:]:
                if len(row) >= 2:  # At minimum need step number/description
                    # Handle various table formats
                    if len(row) >= 3:
                        # Standard format: Step #, Description, Result
                        steps.append(TestStep(
                            number=row[0].strip(),
                            description=row[1].strip(),
                            result=row[2].strip()
                        ))
                    elif len(row) == 2:
                        # Minimal format: Description, Result
                        steps.append(TestStep(
                            number="",
                            description=row[0].strip(),
                            result=row[1].strip()
                        ))
        except Exception as e:
            print(f"  ⚠️  Warning: Error parsing test steps: {e}")
        
        return steps
    
    @staticmethod
    def is_summary_table(table_data: Dict[str, Any]) -> bool:
        """Check if table contains test summary"""
        try:
            if not table_data.get('rows'):
                return False
            
            first_col = [row[0].lower() if row else '' for row in table_data['rows']]
            return any('count of total tests' in cell or 'total tests' in cell for cell in first_col)
        except Exception:
            return False
    
    @staticmethod
    def parse_summary(table_data: Dict[str, Any]) -> Dict[str, str]:
        """Parse test summary from table"""
        summary = {}
        
        try:
            # Check if this is the standard summary table format
            # Header: Count of Total Tests | Count of Completed Tests | % | Count of Passed | Count of Failed | %
            # Data:   21                   | 21                        | 100% | 21            | 0               | 100%
            
            if len(table_data['rows']) >= 2:
                header = table_data['rows'][0]
                data = table_data['rows'][1]
                
                # Try to parse by column position (more reliable)
                if len(data) >= 6:
                    summary['total_tests'] = data[0].strip()
                    summary['completed_tests'] = data[1].strip()
                    summary['percent_completed'] = data[2].strip()
                    summary['passed_tests'] = data[3].strip()
                    summary['failed_tests'] = data[4].strip()
                    summary['percent_passed'] = data[5].strip()
                    
                    # Extract actual scenario count from "Count of Completed Tests"
                    try:
                        completed = int(data[1].strip())
                        summary['actual_scenario_count'] = str(completed)
                        print(f"  ✅ Detected {completed} actual test scenarios from Unit Test Summary")
                    except ValueError:
                        pass
                
                # Fallback: parse by key-value pairs
                else:
                    for row in table_data['rows']:
                        if len(row) >= 2:
                            key = row[0].lower()
                            value = row[1]
                            
                            if 'total tests' in key or 'count of total' in key:
                                summary['total_tests'] = value
                            elif 'completed' in key and '%' not in key:
                                summary['completed_tests'] = value
                                # Extract count
                                try:
                                    completed = int(value.strip())
                                    summary['actual_scenario_count'] = str(completed)
                                    print(f"  ✅ Detected {completed} actual test scenarios from Unit Test Summary")
                                except ValueError:
                                    pass
                            elif 'passed' in key and '%' not in key:
                                summary['passed_tests'] = value
                            elif 'failed' in key and '%' not in key:
                                summary['failed_tests'] = value
                            elif 'passed' in key and '%' in key:
                                summary['percent_passed'] = value
        except Exception as e:
            print(f"  ⚠️  Warning: Error parsing summary: {e}")
        
        return summary


class TES070Analyzer:
    """Main analyzer for TES-070 documents"""
    
    def __init__(self, docx_path: str):
        self.docx_path = docx_path
        self.doc = Document(docx_path)
        self.analysis = {
            'document_info': {},
            'test_summary': {},
            'scenarios': [],
            'document_flow': []
        }
    
    def analyze_complete(self) -> Dict[str, Any]:
        """Perform complete analysis of TES-070 document"""
        print("🔍 Analyzing TES-070 document...")
        
        try:
            self._extract_document_info()
            self._extract_test_summary()
            self._extract_sequential_content()
            self._build_scenarios_from_flow()
        except Exception as e:
            print(f"❌ Error during analysis: {e}")
            raise
        
        return self.analysis
    
    def _extract_document_info(self):
        """Extract document metadata"""
        print("📄 Extracting document info...")
        
        try:
            # Get title from paragraphs
            title = self._find_title()
            
            # Extract from tables
            metadata = self._extract_metadata_from_tables()
            
            self.analysis['document_info'] = {
                'title': title,
                'file_path': self.docx_path,
                **metadata
            }
        except Exception as e:
            print(f"  ⚠️  Warning: Error extracting document info: {e}")
            self.analysis['document_info'] = {
                'title': 'Unknown',
                'file_path': self.docx_path
            }
    
    def _find_title(self) -> str:
        """Find document title"""
        for para in self.doc.paragraphs[:15]:
            text = para.text.strip()
            if text and len(text) > 10 and ('INT_' in text or 'TES-070' in text):
                return text
        return "Unknown"
    
    def _extract_metadata_from_tables(self) -> Dict[str, str]:
        """Extract metadata from document property tables"""
        metadata = {}
        
        for table in self.doc.tables:
            try:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    if len(cells) >= 2:
                        key = cells[0].lower()
                        value = cells[1]
                        
                        if 'author' in key:
                            metadata['author'] = value
                        elif 'version' in key:
                            metadata['version'] = value
                        elif 'date' in key or 'updated' in key:
                            metadata['date'] = value
            except Exception:
                continue
        
        return metadata
    
    def _extract_test_summary(self):
        """Extract test summary statistics"""
        print("📊 Extracting test summary...")
        
        try:
            for table in self.doc.tables:
                table_data = TableParser.parse(table)
                
                if TableParser.is_summary_table(table_data):
                    self.analysis['test_summary'] = TableParser.parse_summary(table_data)
                    break
        except Exception as e:
            print(f"  ⚠️  Warning: Error extracting test summary: {e}")
    
    def _extract_sequential_content(self):
        """Extract all content in document order"""
        print("📄 Extracting content in document order...")
        
        try:
            for element in self.doc.element.body:
                if isinstance(element, CT_P):
                    self._process_paragraph(element)
                elif isinstance(element, CT_Tbl):
                    self._process_table(element)
        except Exception as e:
            print(f"  ⚠️  Warning: Error extracting sequential content: {e}")
    
    def _process_paragraph(self, element):
        """Process a paragraph element"""
        try:
            para = Paragraph(element, self.doc)
            text = para.text.strip()
            
            # Add text if present
            if text:
                self.analysis['document_flow'].append({
                    'type': ContentType.TEXT.value,
                    'content': text,
                    'style': para.style.name if para.style else 'Normal'
                })
        
        except Exception as e:
            print(f"  ⚠️  Warning: Error processing paragraph: {e}")
    
    def _process_table(self, element):
        """Process a table element"""
        try:
            table = Table(element, self.doc)
            table_data = TableParser.parse(table)
            
            # Add table
            self.analysis['document_flow'].append({
                'type': ContentType.TABLE.value,
                'content': table_data
            })
        
        except Exception as e:
            print(f"  ⚠️  Warning: Error processing table: {e}")
    
    def _build_scenarios_from_flow(self):
        """Build scenario structure from document flow"""
        print("🎯 Building test scenarios from content flow...")
        
        current_scenario = None
        current_section = None
        
        try:
            for item in self.analysis['document_flow']:
                if item['type'] == ContentType.TEXT.value:
                    current_scenario, current_section = self._process_text_item(
                        item, current_scenario, current_section
                    )
                
                elif item['type'] == ContentType.TABLE.value and current_scenario:
                    self._process_table_item(item, current_scenario)
            
            # Add last scenario
            if current_scenario:
                self.analysis['scenarios'].append(current_scenario)
            
            # Filter out duplicate/incomplete scenarios (numbered headings vs full scenarios)
            self._deduplicate_scenarios()
        
        except Exception as e:
            print(f"  ⚠️  Warning: Error building scenarios: {e}")
    
    def _deduplicate_scenarios(self):
        """Remove duplicate scenarios (keep ones with more content)"""
        # Group scenarios by similar titles
        unique_scenarios = []
        seen_titles = set()
        
        for scenario in self.analysis['scenarios']:
            # Normalize title for comparison (remove numbers, extra spaces)
            normalized = scenario['title'].lower()
            normalized = ''.join(c for c in normalized if c.isalpha() or c.isspace())
            normalized = ' '.join(normalized.split())
            
            # Check if we've seen a similar title
            if normalized not in seen_titles:
                seen_titles.add(normalized)
                unique_scenarios.append(scenario)
            else:
                # If we've seen it, keep the one with more content
                existing_idx = next(i for i, s in enumerate(unique_scenarios) 
                                  if normalized in ''.join(c for c in s['title'].lower() 
                                  if c.isalpha() or c.isspace()))
                
                existing = unique_scenarios[existing_idx]
                # Keep the one with more test steps
                if len(scenario['test_steps']) > len(existing['test_steps']):
                    unique_scenarios[existing_idx] = scenario
        
        self.analysis['scenarios'] = unique_scenarios
    
    def _process_text_item(self, item: Dict, current_scenario: Optional[Dict], 
                          current_section: Optional[str]) -> tuple:
        """Process a text item in document flow"""
        text = item['content']
        
        # Detect scenario start (handle both "Scenario:" and numbered "3.1 Scenario:")
        if 'scenario:' in text.lower() and (text.strip().startswith('Scenario:') or 
                                            any(text.strip().startswith(f'{i}.') for i in range(1, 10))):
            if current_scenario:
                self.analysis['scenarios'].append(current_scenario)
            
            current_scenario = {
                'title': text,
                'description': '',
                'test_steps': [],
                'results': [],
                'content_flow': [item]
            }
            current_section = 'scenario'
        
        elif current_scenario:
            # Track section headers (handle both plain and numbered formats)
            text_lower = text.lower().strip()
            
            if text_lower == 'description' or text_lower.endswith('description'):
                current_section = 'description'
            elif text_lower == 'test steps' or text_lower.endswith('test steps'):
                current_section = 'test_steps'
            elif text_lower == 'results' or text_lower.endswith('results'):
                current_section = 'results'
            elif text and current_section == 'description' and not text_lower.startswith(('test steps', 'results', 'scenario')):
                current_scenario['description'] += text + ' '
            elif text and current_section == 'results' and not text_lower.startswith(('scenario', 'test steps')):
                current_scenario['results'].append(text)
            
            current_scenario['content_flow'].append(item)
        
        return current_scenario, current_section
    
    def _process_table_item(self, item: Dict, current_scenario: Dict):
        """Process a table item in document flow"""
        table_data = item['content']
        
        if TableParser.is_test_steps_table(table_data):
            steps = TableParser.parse_test_steps(table_data)
            current_scenario['test_steps'] = [asdict(step) for step in steps]
        
        current_scenario['content_flow'].append(item)
    
    def save_analysis(self, output_path: Optional[str] = None):
        """Save analysis to JSON file"""
        try:
            if output_path is None:
                source_path = Path(self.docx_path)
                output_path = source_path.parent / (source_path.stem + "_analysis.json")
            
            with open(output_path, 'w', encoding='utf-8') as f:
                json.dump(self.analysis, f, indent=2, ensure_ascii=False)
            
            print(f"\n💾 Analysis saved to: {output_path}")
        except Exception as e:
            print(f"❌ Error saving analysis: {e}")
    
    def print_summary(self):
        """Print analysis summary"""
        print("\n" + "="*60)
        print("TES-070 ANALYSIS SUMMARY")
        print("="*60)
        
        # Document info
        doc_info = self.analysis['document_info']
        print(f"\n📄 Document: {doc_info.get('title', 'N/A')}")
        print(f"   Author: {doc_info.get('author', 'N/A')}")
        print(f"   Version: {doc_info.get('version', 'N/A')}")
        
        # Test summary
        if self.analysis['test_summary']:
            print(f"\n📊 Test Summary:")
            summary = self.analysis['test_summary']
            print(f"   Total Tests: {summary.get('total_tests', 'N/A')}")
            print(f"   Passed: {summary.get('passed_tests', 'N/A')}")
            print(f"   Failed: {summary.get('failed_tests', 'N/A')}")
            print(f"   Pass Rate: {summary.get('percent_passed', 'N/A')}")
            
            # Show actual scenario count if detected
            if 'actual_scenario_count' in summary:
                print(f"\n✅ ACTUAL SCENARIO COUNT: {summary['actual_scenario_count']}")
                print(f"   (Extracted from Unit Test Summary - Section 2.1)")
        
        # Scenarios
        detected_count = len(self.analysis['scenarios'])
        actual_count = self.analysis['test_summary'].get('actual_scenario_count', 'Unknown')
        
        print(f"\n🎯 Scenarios Detected in Document: {detected_count}")
        if actual_count != 'Unknown':
            print(f"   ⚠️  Note: Unit Test Summary indicates {actual_count} actual scenarios")
            if str(detected_count) != str(actual_count):
                print(f"   ⚠️  Mismatch detected - document may have TOC entries + detailed scenarios")
        for idx, scenario in enumerate(self.analysis['scenarios'], 1):
            print(f"   {idx}. {scenario['title']}")
            print(f"      Steps: {len(scenario['test_steps'])}")
        
        # Document flow
        print(f"\n📋 Document Flow: {len(self.analysis['document_flow'])} items")
        flow_types = {}
        for item in self.analysis['document_flow']:
            item_type = item['type']
            flow_types[item_type] = flow_types.get(item_type, 0) + 1
        
        for flow_type, count in flow_types.items():
            print(f"   {flow_type}: {count}")
        
        print("\n" + "="*60)


def analyze_tes070(docx_path: str, save_json: bool = True) -> Dict[str, Any]:
    """
    Analyze a TES-070 document completely
    
    Args:
        docx_path: Path to TES-070 .docx file
        save_json: Whether to save analysis to JSON file
    
    Returns:
        Complete analysis dictionary
    """
    # Normalize path to handle special characters
    docx_path = os.path.normpath(docx_path)
    
    # Verify file exists
    if not os.path.exists(docx_path):
        # Try with Path for better Unicode handling
        from pathlib import Path
        path_obj = Path(docx_path)
        if not path_obj.exists():
            raise FileNotFoundError(f"File not found: {docx_path}")
        docx_path = str(path_obj.absolute())
    
    analyzer = TES070Analyzer(docx_path)
    analysis = analyzer.analyze_complete()
    
    analyzer.print_summary()
    
    if save_json:
        analyzer.save_analysis()
    
    return analysis


if __name__ == "__main__":
    import sys
    
    if len(sys.argv) < 2:
        print("Usage: python tes070_analyzer.py <path_to_tes070.docx>")
        print("\nExample:")
        print('  python tes070_analyzer.py "TES-070/document.docx"')
        sys.exit(1)
    
    docx_file = sys.argv[1]
    
    if os.path.exists(docx_file):
        try:
            analyze_tes070(docx_file)
        except Exception as e:
            print(f"\n❌ Analysis failed: {e}")
            sys.exit(1)
    else:
        print(f"❌ File not found: {docx_file}")
        sys.exit(1)
