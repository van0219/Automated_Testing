"""
TES-070 Document Generator - Create test results documents with proper formatting

Purpose: Generate TES-070 documents from test execution data
Use when: You need to create official test results documentation

Features:
- Template-based generation with exact formatting from samples
- Automatic styling (Arial fonts, Infor colors, proper spacing)
- Screenshot embedding in correct locations
- Table of contents generation
- Summary statistics calculation
- Follows TES-070 standards from steering file
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path
from typing import Dict, List, Optional, Any
from dataclasses import dataclass
from datetime import datetime
import os


@dataclass
class TestStep:
    """Single test step"""
    number: str
    description: str
    result: str  # PASS, FAIL, or empty


@dataclass
class TestScenario:
    """Complete test scenario"""
    title: str
    description: str
    test_steps: List[TestStep]
    results: List[str]
    screenshots: List[str]  # Paths to screenshot files


@dataclass
class TES070Data:
    """Complete TES-070 document data"""
    # Document metadata
    client_name: str
    interface_id: str
    interface_name: str
    author: str
    version: str = "1.0"
    date: str = None
    
    # Test summary
    total_tests: int = 0
    completed_tests: int = 0
    passed_tests: int = 0
    failed_tests: int = 0
    
    # Prerequisites
    environment: str = ""
    user_roles: List[str] = None
    test_data_requirements: str = ""
    configuration_prerequisites: str = ""
    
    # Test scenarios
    scenarios: List[TestScenario] = None
    
    def __post_init__(self):
        if self.date is None:
            self.date = datetime.now().strftime("%m/%d/%Y")
        if self.user_roles is None:
            self.user_roles = []
        if self.scenarios is None:
            self.scenarios = []
        
        # Calculate summary if not provided
        if self.total_tests == 0:
            self.total_tests = len(self.scenarios)
        if self.completed_tests == 0:
            self.completed_tests = len([s for s in self.scenarios if any(step.result for step in s.test_steps)])
        if self.passed_tests == 0:
            self.passed_tests = len([s for s in self.scenarios if all(step.result == 'PASS' for step in s.test_steps if step.result)])
        if self.failed_tests == 0:
            self.failed_tests = self.completed_tests - self.passed_tests
    
    @property
    def percent_completed(self) -> str:
        if self.total_tests == 0:
            return "0%"
        return f"{(self.completed_tests / self.total_tests * 100):.0f}%"
    
    @property
    def percent_passed(self) -> str:
        if self.completed_tests == 0:
            return "0%"
        return f"{(self.passed_tests / self.completed_tests * 100):.0f}%"


class TES070StyleManager:
    """Manages document styles matching TES-070 standards"""
    
    # Infor Blue color
    INFOR_BLUE = RGBColor(19, 163, 247)  # #13A3F7
    BLACK = RGBColor(0, 0, 0)
    WHITE = RGBColor(255, 255, 255)
    GRAY = RGBColor(242, 242, 242)
    
    @staticmethod
    def setup_styles(doc: Document):
        """Setup all required styles for TES-070 document"""
        styles = doc.styles
        
        # Heading 1: 14pt bold, ALL CAPS
        h1 = styles['Heading 1']
        h1.font.name = 'Arial'
        h1.font.size = Pt(14)
        h1.font.bold = True
        h1.font.color.rgb = TES070StyleManager.BLACK
        
        # Heading 2: 14pt bold, Title Case
        h2 = styles['Heading 2']
        h2.font.name = 'Arial'
        h2.font.size = Pt(14)
        h2.font.bold = True
        h2.font.color.rgb = TES070StyleManager.BLACK
        
        # Heading 3: 12pt bold, ALL CAPS
        h3 = styles['Heading 3']
        h3.font.name = 'Arial'
        h3.font.size = Pt(12)
        h3.font.bold = True
        h3.font.color.rgb = TES070StyleManager.BLACK
        
        # Heading 4: 12pt bold, Title Case
        h4 = styles['Heading 4']
        h4.font.name = 'Arial'
        h4.font.size = Pt(12)
        h4.font.bold = True
        h4.font.color.rgb = TES070StyleManager.BLACK
        
        # Normal: 11pt
        normal = styles['Normal']
        normal.font.name = 'Arial'
        normal.font.size = Pt(11)
        normal.font.color.rgb = TES070StyleManager.BLACK
    
    @staticmethod
    def format_table(table, header_row=True):
        """Apply TES-070 table formatting"""
        # Set table style
        table.style = 'Table Grid'
        
        # Format header row
        if header_row and len(table.rows) > 0:
            for cell in table.rows[0].cells:
                # Black background, white text, bold
                cell_shading = cell._element.get_or_add_tcPr()
                shading_elm = OxmlElement('w:shd')
                shading_elm.set(qn('w:fill'), '000000')
                cell_shading.append(shading_elm)
                
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.color.rgb = TES070StyleManager.WHITE
                        run.font.bold = True
                        run.font.size = Pt(10)
                        run.font.name = 'Arial'
        
        # Format body rows
        for i, row in enumerate(table.rows[1:] if header_row else table.rows, start=1):
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(10)
                        run.font.name = 'Arial'


class TES070Generator:
    """Main generator for TES-070 documents"""
    
    def __init__(self, data: TES070Data):
        self.data = data
        self.doc = Document()
        TES070StyleManager.setup_styles(self.doc)
    
    def generate(self, output_path: str) -> str:
        """Generate complete TES-070 document"""
        print("📝 Generating TES-070 document...")
        
        # Generate all sections
        self._add_title_page()
        self._add_page_break()
        self._add_document_control()
        self._add_page_break()
        self._add_table_of_contents()
        self._add_page_break()
        self._add_test_summary()
        self._add_prerequisites()
        
        # Add each scenario
        for i, scenario in enumerate(self.data.scenarios, 1):
            self._add_scenario(i, scenario)
        
        # Save document
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(str(output_path))
        
        print(f"✅ TES-070 document generated: {output_path}")
        return str(output_path)
    
    def _add_title_page(self):
        """Add title page with Infor logo"""
        print("  📄 Adding title page...")
        
        # Add Infor logo at the top left (small size like in samples)
        logo_path = Path("Infor_Logo.jpg")
        if logo_path.exists():
            try:
                logo_para = self.doc.add_paragraph()
                logo_para.alignment = WD_ALIGN_PARAGRAPH.LEFT  # Left aligned, not centered
                logo_run = logo_para.add_run()
                logo_run.add_picture(str(logo_path), width=Inches(0.6))  # Much smaller - 0.6 inches
                self.doc.add_paragraph()  # Add spacing after logo
                print("    ✅ Infor logo added")
            except Exception as e:
                print(f"    ⚠️  Warning: Could not add Infor logo: {e}")
        else:
            print("    ⚠️  Warning: Infor logo not found at Infor_Logo.jpg")
        
        # Title
        title = self.doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run("TES-070 Custom Interface Unit Test Results")
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.name = 'Arial'
        
        # Client name
        client = self.doc.add_paragraph()
        client.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = client.add_run(self.data.client_name)
        run.font.size = Pt(14)
        run.font.name = 'Arial'
        
        # Interface ID and Name
        interface = self.doc.add_paragraph()
        interface.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = interface.add_run(f"{self.data.interface_id} {self.data.interface_name}")
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.name = 'Arial'
        
        # Add spacing
        self.doc.add_paragraph()
        
        # Document properties table
        table = self.doc.add_table(rows=5, cols=2)
        table.style = 'Table Grid'
        
        properties = [
            ("Author", self.data.author),
            ("Creation Date", self.data.date),
            ("Last Updated", self.data.date),
            ("Version", self.data.version),
            ("Document Reference", "")
        ]
        
        for i, (key, value) in enumerate(properties):
            table.rows[i].cells[0].text = key
            table.rows[i].cells[1].text = value
    
    def _add_document_control(self):
        """Add document control section"""
        print("  📋 Adding document control...")
        
        heading = self.doc.add_heading("1 Document Control", level=1)
        
        # Change Record
        self.doc.add_heading("1.1 Change Record", level=2)
        table = self.doc.add_table(rows=2, cols=4)
        TES070StyleManager.format_table(table)
        
        headers = ["Date", "Author", "Version", "Change Description"]
        for i, header in enumerate(headers):
            table.rows[0].cells[i].text = header
        
        table.rows[1].cells[0].text = self.data.date
        table.rows[1].cells[1].text = self.data.author
        table.rows[1].cells[2].text = self.data.version
        table.rows[1].cells[3].text = "Initial version"
        
        # Reviewers
        self.doc.add_heading("1.2 Reviewers", level=2)
        table = self.doc.add_table(rows=2, cols=2)
        TES070StyleManager.format_table(table)
        
        table.rows[0].cells[0].text = "Name"
        table.rows[0].cells[1].text = "Position"
        table.rows[1].cells[0].text = ""
        table.rows[1].cells[1].text = ""
    
    def _add_table_of_contents(self):
        """Add table of contents with proper Word field"""
        print("  📑 Adding table of contents...")
        
        heading = self.doc.add_heading("Contents", level=1)
        
        # Add TOC field that Word can update
        paragraph = self.doc.add_paragraph()
        run = paragraph.add_run()
        
        # Create TOC field code with proper structure
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
        
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')
        
        # Add placeholder text between separate and end
        t = OxmlElement('w:t')
        t.text = 'Right-click and select "Update Field" to generate table of contents'
        
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'end')
        
        r_element = run._element
        r_element.append(fldChar1)
        r_element.append(instrText)
        r_element.append(fldChar2)
        r_element.append(t)
        r_element.append(fldChar3)
        
        # Style the paragraph
        run.font.italic = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(128, 128, 128)
    
    def _add_test_summary(self):
        """Add test summary section"""
        print("  📊 Adding test summary...")
        
        self.doc.add_heading("2 Custom Interface Unit Test Summary Results", level=1)
        
        # Unit Test Summary
        self.doc.add_heading("2.1 Unit Test Summary", level=2)
        
        p = self.doc.add_paragraph(
            "The Custom Interface Unit Test Results consolidate the findings from "
            "the unit testing performed on the custom interface."
        )
        
        # Summary table
        table = self.doc.add_table(rows=2, cols=6)
        TES070StyleManager.format_table(table)
        
        headers = [
            "Count of Total Tests",
            "Count of Completed Tests",
            "% of Completed Tests",
            "Count of Passed Tests",
            "Count of Failed Tests",
            "% of Passed Tests"
        ]
        
        for i, header in enumerate(headers):
            table.rows[0].cells[i].text = header
        
        values = [
            str(self.data.total_tests),
            str(self.data.completed_tests),
            self.data.percent_completed,
            str(self.data.passed_tests),
            str(self.data.failed_tests),
            self.data.percent_passed
        ]
        
        for i, value in enumerate(values):
            table.rows[1].cells[i].text = value
        
        # Test Scenarios
        self.doc.add_heading("2.2 Test Scenarios", level=2)
        
        table = self.doc.add_table(rows=len(self.data.scenarios) + 1, cols=3)
        TES070StyleManager.format_table(table)
        
        table.rows[0].cells[0].text = "Scenario #"
        table.rows[0].cells[1].text = "Test Scenario"
        table.rows[0].cells[2].text = "Infor Test Result"
        
        for i, scenario in enumerate(self.data.scenarios, 1):
            table.rows[i].cells[0].text = str(i)
            table.rows[i].cells[1].text = scenario.title
            # Determine result
            if all(step.result == 'PASS' for step in scenario.test_steps if step.result):
                result = "PASS"
            elif any(step.result == 'FAIL' for step in scenario.test_steps):
                result = "FAIL"
            else:
                result = "In Progress"
            table.rows[i].cells[2].text = result
    
    def _add_prerequisites(self):
        """Add prerequisites section"""
        print("  ⚙️  Adding prerequisites...")
        
        self.doc.add_heading("3 Custom Interface Unit Test Detailed Results", level=1)
        
        self.doc.add_heading("Prerequisites for testing:", level=3)
        
        if self.data.environment:
            p = self.doc.add_paragraph()
            p.add_run("Environment: ").bold = True
            p.add_run(self.data.environment)
        
        if self.data.user_roles:
            p = self.doc.add_paragraph()
            p.add_run("User Roles: ").bold = True
            p.add_run(", ".join(self.data.user_roles))
        
        if self.data.test_data_requirements:
            p = self.doc.add_paragraph()
            p.add_run("Test Data: ").bold = True
            p.add_run(self.data.test_data_requirements)
        
        if self.data.configuration_prerequisites:
            p = self.doc.add_paragraph()
            p.add_run("Configuration: ").bold = True
            p.add_run(self.data.configuration_prerequisites)
    
    def _add_scenario(self, scenario_num: int, scenario: TestScenario):
        """Add a complete test scenario"""
        print(f"  🎯 Adding scenario {scenario_num}: {scenario.title[:50]}...")
        
        # Scenario heading
        self.doc.add_heading(f"3.{scenario_num} {scenario.title}", level=2)
        
        # Description
        self.doc.add_heading(f"3.{scenario_num}.1 Description", level=3)
        self.doc.add_paragraph(scenario.description)
        
        # Test Steps
        self.doc.add_heading(f"3.{scenario_num}.2 Test Steps", level=3)
        
        if scenario.test_steps:
            table = self.doc.add_table(rows=len(scenario.test_steps) + 1, cols=3)
            TES070StyleManager.format_table(table)
            
            table.rows[0].cells[0].text = "No."
            table.rows[0].cells[1].text = "Steps"
            table.rows[0].cells[2].text = "Result"
            
            for i, step in enumerate(scenario.test_steps, 1):
                table.rows[i].cells[0].text = step.number or str(i)
                table.rows[i].cells[1].text = step.description
                table.rows[i].cells[2].text = step.result
        
        # Embed screenshots
        if scenario.screenshots:
            print(f"    📸 Embedding {len(scenario.screenshots)} screenshots...")
            for screenshot_path in scenario.screenshots:
                if Path(screenshot_path).exists():
                    try:
                        self.doc.add_picture(screenshot_path, width=Inches(6))
                        self.doc.add_paragraph()  # Add spacing
                    except Exception as e:
                        print(f"    ⚠️  Warning: Could not embed {screenshot_path}: {e}")
                else:
                    print(f"    ⚠️  Warning: Screenshot not found: {screenshot_path}")
        
        # Results
        if scenario.results:
            self.doc.add_heading(f"3.{scenario_num}.3 Results", level=3)
            for result in scenario.results:
                self.doc.add_paragraph(result, style='List Bullet')
    
    def _add_page_break(self):
        """Add page break"""
        self.doc.add_page_break()


def generate_tes070(data: TES070Data, output_path: str) -> str:
    """
    Generate a TES-070 document
    
    Args:
        data: TES070Data object with all document information
        output_path: Path where to save the generated document
    
    Returns:
        Path to generated document
    
    Example:
        >>> data = TES070Data(
        ...     client_name="State of New Hampshire",
        ...     interface_id="INT_FIN_013",
        ...     interface_name="GL Transaction Interface",
        ...     author="Test Engineer",
        ...     environment="ACUITY_TST",
        ...     user_roles=["Process Server Administrator", "Financials Processor"],
        ...     scenarios=[...]
        ... )
        >>> generate_tes070(data, "TES-070/Generated_TES070s/test_results.docx")
    """
    generator = TES070Generator(data)
    return generator.generate(output_path)


if __name__ == "__main__":
    # Example usage
    print("TES-070 Generator")
    print("=" * 60)
    print("\nThis tool generates TES-070 test results documents.")
    print("\nUsage:")
    print("  from tes070_generator import generate_tes070, TES070Data, TestScenario, TestStep")
    print("\nSee docstrings for detailed usage examples.")
